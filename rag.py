"""
RAG (Retrieval-Augmented Generation) System
Processes documents from classification_documents folder and stores them in Pinecone
Each document gets its own namespace based on the document name
Uses OpenAI ada-002 embedding model
"""

import streamlit as st
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib

# Third-party imports
from pinecone import Pinecone, ServerlessSpec
from openai import OpenAI
import docx2txt
from docx import Document

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentRAG:
    """RAG system for processing and storing documents in Pinecone"""
    
    def __init__(self):
        """Initialize the RAG system with API keys and clients"""
        try:
            # Initialize OpenAI client
            self.openai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            
            # Initialize Pinecone
            self.pc = Pinecone(api_key=st.secrets["PINECONE_API_KEY"])
            self.index_name = st.secrets["PINECONE_INDEX"]
            self.pinecone_env = st.secrets["PINECONE_ENV"]
            
            # Get or create index
            self.index = self._get_or_create_index()
            
            # Embedding model configuration
            self.embedding_model = "text-embedding-ada-002"
            self.embedding_dimension = 1024  # Match Pinecone index dimension
            self.chunk_size = 1000  # Characters per chunk
            self.chunk_overlap = 200  # Overlap between chunks
            
            logger.info("DocumentRAG initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing DocumentRAG: {e}")
            raise
    
    def _get_or_create_index(self):
        """Get existing index or create new one if it doesn't exist"""
        try:
            # Check if index exists
            existing_indexes = [index.name for index in self.pc.list_indexes()]
            
            if self.index_name not in existing_indexes:
                logger.info(f"Creating new index: {self.index_name}")
                self.pc.create_index(
                    name=self.index_name,
                    dimension=self.embedding_dimension,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region=self.pinecone_env
                    )
                )
                logger.info(f"Index {self.index_name} created successfully")
            else:
                logger.info(f"Using existing index: {self.index_name}")
            
            return self.pc.Index(self.index_name)
            
        except Exception as e:
            logger.error(f"Error with Pinecone index: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from .docx file"""
        try:
            # Method 1: Using docx2txt (simpler)
            text = docx2txt.process(file_path)
            
            if not text.strip():
                # Method 2: Using python-docx (more detailed)
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks"""
        if chunk_size is None:
            chunk_size = self.chunk_size
        if overlap is None:
            overlap = self.chunk_overlap
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence end
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                sentence_ends = ['.', '!', '?', '\\n']
                best_break = end
                
                for i in range(max(0, end - 100), min(len(text), end + 100)):
                    if text[i] in sentence_ends:
                        best_break = i + 1
                        break
                
                end = best_break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
            # Prevent infinite loop
            if start >= len(text):
                break
        
        return chunks
    
    def get_embedding(self, text: str) -> List[float]:
        """Get embedding for text using OpenAI ada-002"""
        try:
            response = self.openai_client.embeddings.create(
                model=self.embedding_model,
                input=text
            )
            # Truncate embedding to match Pinecone index dimension (1024)
            embedding = response.data[0].embedding
            return embedding[:self.embedding_dimension]
            
        except Exception as e:
            logger.error(f"Error getting embedding: {e}")
            raise
    
    def create_document_id(self, doc_name: str, chunk_index: int) -> str:
        """Create unique ID for document chunk"""
        content = f"{doc_name}_{chunk_index}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def get_namespace_from_filename(self, filename: str) -> str:
        """Extract namespace from filename (remove extension and clean)"""
        # Remove file extension
        name = Path(filename).stem
        # Clean the name to be a valid namespace
        namespace = name.replace(" ", "_").replace("-", "_").lower()
        return namespace
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a single document and return metadata"""
        try:
            filename = os.path.basename(file_path)
            namespace = self.get_namespace_from_filename(filename)
            
            logger.info(f"Processing document: {filename} -> namespace: {namespace}")
            
            # Extract text
            text = self.extract_text_from_docx(file_path)
            if not text:
                logger.warning(f"No text extracted from {filename}")
                return {"status": "failed", "reason": "No text extracted"}
            
            # Chunk text
            chunks = self.chunk_text(text)
            logger.info(f"Created {len(chunks)} chunks for {filename}")
            
            # Process each chunk
            vectors_to_upsert = []
            
            for i, chunk in enumerate(chunks):
                try:
                    # Get embedding
                    embedding = self.get_embedding(chunk)
                    
                    # Create unique ID
                    doc_id = self.create_document_id(filename, i)
                    
                    # Prepare metadata
                    metadata = {
                        "document_name": filename,
                        "chunk_index": i,
                        "chunk_text": chunk,
                        "total_chunks": len(chunks),
                        "file_path": file_path,
                        "namespace": namespace
                    }
                    
                    # Prepare vector for upsert
                    vector = {
                        "id": doc_id,
                        "values": embedding,
                        "metadata": metadata
                    }
                    
                    vectors_to_upsert.append(vector)
                    
                except Exception as e:
                    logger.error(f"Error processing chunk {i} of {filename}: {e}")
                    continue
            
            # Upsert vectors to Pinecone with namespace
            if vectors_to_upsert:
                try:
                    self.index.upsert(
                        vectors=vectors_to_upsert,
                        namespace=namespace
                    )
                    logger.info(f"Successfully upserted {len(vectors_to_upsert)} vectors for {filename}")
                    
                    return {
                        "status": "success",
                        "filename": filename,
                        "namespace": namespace,
                        "chunks_processed": len(vectors_to_upsert),
                        "total_chunks": len(chunks)
                    }
                    
                except Exception as e:
                    logger.error(f"Error upserting vectors for {filename}: {e}")
                    return {"status": "failed", "reason": f"Upsert error: {e}"}
            else:
                return {"status": "failed", "reason": "No valid chunks to upsert"}
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {"status": "failed", "reason": str(e)}
    
    def process_all_documents(self, folder_path: str = "classification_documents") -> List[Dict[str, Any]]:
        """Process all .docx documents in the specified folder"""
        try:
            folder_path_obj = Path(folder_path)
            if not folder_path_obj.exists():
                logger.error(f"Folder {folder_path} does not exist")
                return []
            
            # Find all .docx files
            docx_files = list(folder_path_obj.glob("*.docx"))
            logger.info(f"Found {len(docx_files)} .docx files to process")
            
            results = []
            
            for file_path in docx_files:
                logger.info(f"Processing: {file_path}")
                result = self.process_document(str(file_path))
                results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Error processing documents: {e}")
            return []
    
    def search_documents(self, query: str, namespace: Optional[str] = None, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for relevant document chunks in specific namespace or all namespaces"""
        try:
            # Get query embedding
            query_embedding = self.get_embedding(query)
            
            # Search in Pinecone
            search_kwargs = {
                "vector": query_embedding,
                "top_k": top_k,
                "include_metadata": True
            }
            
            # Add namespace filter if specified
            if namespace and namespace.strip():
                search_kwargs["namespace"] = namespace
                logger.info(f"Searching in namespace: {namespace}")
            else:
                logger.info("Searching across all namespaces")
            
            results = self.index.query(**search_kwargs)
            
            # Format results
            formatted_results = []
            for match in results.matches:
                formatted_results.append({
                    "id": match.id,
                    "score": match.score,
                    "text": match.metadata.get("chunk_text", ""),
                    "document_name": match.metadata.get("document_name", ""),
                    "chunk_index": match.metadata.get("chunk_index", 0),
                    "namespace": match.metadata.get("namespace", "")
                })
            
            logger.info(f"Found {len(formatted_results)} results for query: '{query}'")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
    
    def get_namespaces(self) -> List[str]:
        """Get list of available namespaces (document names)"""
        try:
            # Get index statistics to see available namespaces
            stats = self.index.describe_index_stats()
            
            # Extract namespace names from stats
            if stats.namespaces:
                namespaces = list(stats.namespaces.keys())
                logger.info(f"Found namespaces: {namespaces}")
                return namespaces
            else:
                # If no namespaces found in stats, try to get from local documents
                logger.warning("No namespaces found in index stats, checking local documents...")
                folder_path = Path("classification_documents")
                if folder_path.exists():
                    docx_files = list(folder_path.glob("*.docx"))
                    namespaces = [self.get_namespace_from_filename(file.name) for file in docx_files]
                    logger.info(f"Generated namespaces from local files: {namespaces}")
                    return namespaces
                else:
                    logger.warning("No classification_documents folder found")
                    return []
            
        except Exception as e:
            logger.error(f"Error getting namespaces: {e}")
            # Fallback: return namespaces based on local files
            try:
                folder_path = Path("classification_documents")
                if folder_path.exists():
                    docx_files = list(folder_path.glob("*.docx"))
                    namespaces = [self.get_namespace_from_filename(file.name) for file in docx_files]
                    return namespaces
            except:
                pass
            return []
    
    def delete_namespace(self, namespace: str) -> bool:
        """Delete all vectors in a specific namespace"""
        try:
            self.index.delete(delete_all=True, namespace=namespace)
            logger.info(f"Deleted namespace: {namespace}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting namespace {namespace}: {e}")
            return False
    
    def get_document_content(self, document_name: str, query: str, max_chunks: int = 3) -> str:
        """Get relevant content from a specific document for RAG purposes"""
        try:
            # Get namespace from document name
            namespace = self.get_namespace_from_filename(document_name)
            
            # Search in that specific namespace
            results = self.search_documents(query, namespace=namespace, top_k=max_chunks)
            
            # Combine the most relevant chunks
            combined_content = ""
            for i, result in enumerate(results):
                combined_content += f"[Chunk {i+1}]\n{result['text']}\n\n"
            
            return combined_content.strip()
            
        except Exception as e:
            logger.error(f"Error getting document content: {e}")
            return ""


# Helper function for easy integration with other parts of the application
def search_classification_docs(query: str, document_section: Optional[str] = None) -> str:
    """
    Helper function to search classification documents and return relevant content
    
    Args:
        query: Search query
        document_section: Specific document section to search in (optional)
    
    Returns:
        Combined relevant content from documents
    """
    try:
        rag = DocumentRAG()
        
        # Determine namespace if document_section is provided
        namespace = None
        if document_section:
            # Convert section name to namespace format
            namespace = document_section.lower().replace(" ", "_").replace("-", "_")
        
        # Search documents
        results = rag.search_documents(query, namespace=namespace, top_k=3)
        
        # Combine results into readable format
        combined_content = ""
        for result in results:
            combined_content += f"{result['text']}\n\n"
        
        return combined_content.strip()
        
    except Exception as e:
        logger.error(f"Error in search_classification_docs: {e}")
        return ""


def main():
    """Main function for testing the RAG system"""
    st.title("Document RAG System")
    
    try:
        # Initialize RAG system
        rag = DocumentRAG()
        
        st.success("RAG system initialized successfully!")
        
        # Process documents button
        if st.button("Process All Documents"):
            with st.spinner("Processing documents..."):
                results = rag.process_all_documents()
                
                st.subheader("Processing Results:")
                for result in results:
                    if result["status"] == "success":
                        st.success(f"✅ {result['filename']}: {result['chunks_processed']} chunks processed")
                    else:
                        st.error(f"❌ {result.get('filename', 'Unknown')}: {result.get('reason', 'Unknown error')}")
        
        # Search interface
        st.subheader("Search Documents")
        query = st.text_input("Enter your search query:")
        
        # Get available namespaces
        with st.spinner("Loading available documents..."):
            namespaces = rag.get_namespaces()
        
        # Create namespace options with better labels
        namespace_options = ["All documents"]
        namespace_mapping: Dict[str, Optional[str]] = {"All documents": None}
        
        for ns in namespaces:
            # Convert namespace back to readable format
            readable_name = ns.replace("_", " ").title()
            namespace_options.append(readable_name)
            namespace_mapping[readable_name] = ns
        
        selected_option = st.selectbox(
            "Select document to search in:", 
            options=namespace_options,
            help="Choose a specific document or search across all documents"
        )
        
        # Show current selection info
        selected_namespace: Optional[str] = None
        if selected_option and selected_option != "All documents":
            selected_namespace = namespace_mapping.get(selected_option)
        if selected_namespace:
            st.info(f"🔍 Searching in: **{selected_option}** (namespace: `{selected_namespace}`)")
        else:
            st.info("🔍 Searching across **all documents**")
        
        # Search button and results
        if st.button("Search", type="primary") and query:
            with st.spinner("Searching documents..."):
                # Pass the selected namespace (None for all documents)
                results = rag.search_documents(query, namespace=selected_namespace)
                
                if results:
                    st.subheader(f"Found {len(results)} relevant chunks:")
                    
                    for i, result in enumerate(results):
                        # Create expandable result card
                        score_color = "🟢" if result['score'] > 0.8 else "🟡" if result['score'] > 0.6 else "🔴"
                        
                        with st.expander(
                            f"{score_color} Result {i+1}: {result['document_name']} (Similarity: {result['score']:.3f})",
                            expanded=(i < 3)  # Expand first 3 results
                        ):
                            # Show the relevant text
                            st.markdown("**Content:**")
                            st.write(result['text'])
                            
                            # Show metadata
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.caption(f"📄 **Document:** {result['document_name']}")
                            with col2:
                                st.caption(f"📝 **Chunk:** {result['chunk_index'] + 1}")
                            with col3:
                                st.caption(f"🏷️ **Namespace:** {result['namespace']}")
                else:
                    st.warning("No relevant documents found for your query.")
                    st.info("Try:")
                    st.write("- Using different keywords")
                    st.write("- Searching in a different document")
                    st.write("- Making sure documents are processed first")
        
        # Show namespace statistics
        if namespaces:
            with st.expander("📊 Document Statistics"):
                st.write(f"**Total documents processed:** {len(namespaces)}")
                for i, ns in enumerate(namespaces, 1):
                    readable_name = ns.replace("_", " ").title()
                    st.write(f"{i}. {readable_name} (`{ns}`)")
        
    except Exception as e:
        st.error(f"Error initializing RAG system: {e}")
        st.info("Make sure your Pinecone and OpenAI credentials are correctly set in `.streamlit/secrets.toml`")


if __name__ == "__main__":
    main()
